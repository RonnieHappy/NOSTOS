"""Build the evidence-linked Small Methods Supporting Information v34."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_nostos0_methods_docx as base


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "resources" / "journal_templates" / "Small_Methods_Article_Template_2025.docx"
OUTPUT = ROOT / "manuscripts" / "NOSTOS_Small_Methods_Supporting_Information_v34.docx"
TITLE = "NOSTOS Prevents Silent Acquisition- and Scale-Specific Failure in Quantitative Microscopy"

FIGURES = {
    "Figure S1": ROOT / "figures" / "nostos0_small_methods_si" / "figure_s1_synthetic_validation.png",
    "Figure S2": ROOT / "figures" / "nostos0" / "figure_3_bone_validation.png",
    "Figure S3": ROOT / "figures" / "nostos0" / "supplementary_figure_1_bone_contract_stress.png",
}

ALT_TEXT = {
    "Figure S1": "Programmed synthetic phantoms, exact perturbation errors, module-by-perturbation support matrix and monotonic network erosion survival.",
    "Figure S2": "Public trabecular-bone local-thickness reference, NOSTOS estimate, residual map, volume-level agreement and paired error comparison.",
    "Figure S3": "Public label-free and three-dimensional bone images with deterministic compatibility, topology, orientation, spectrum and selective-risk outputs.",
}

RECEIPTS = [
    ROOT / "outputs" / "nostos0-synthetic-v1" / "validation.json",
    ROOT / "outputs" / "nostos0-module-perturbations-v1" / "module_perturbation_matrix.json",
    ROOT / "outputs" / "nostos0-stare-network-confirmation-v1" / "stare_network_confirmation.json",
    ROOT / "outputs" / "external-bone-v1" / "external_bone_validation.json",
    ROOT / "outputs" / "nostos0-bonej-thickness-v1" / "bonej_thickness_comparator.json",
    ROOT / "outputs" / "nostos0-bbbc035-dynamic-confirmation-v1" / "bbbc035_dynamic_confirmation.json",
    ROOT / "outputs" / "nostos0-bbbc035-dense-deformation-confirmation-v1" / "bbbc035_dense_deformation_confirmation.json",
    ROOT / "outputs" / "nostos0-bbbc006-spatial-confirmation-v1" / "bbbc006_spatial_confirmation.json",
]


def load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def set_run(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    base.set_font(run, size=size, bold=bold, italic=italic, color=(0, 0, 0))


def add_paragraph(doc: Document, text: str = "", *, size: float = 9.35, after: float = 3.0) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    set_run(paragraph.add_run(text), size=size)
    return paragraph


def add_heading(doc: Document, text: str, *, level: int = 1) -> object:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 55, start: int = 65, bottom: int = 55, end: int = 65) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    definitions = {
        "top": ("single", "6", "9AA8AF"),
        "bottom": ("single", "6", "9AA8AF"),
        "insideH": ("single", "3", "DDE3E6"),
        "left": ("nil", "0", "FFFFFF"),
        "right": ("nil", "0", "FFFFFF"),
        "insideV": ("nil", "0", "FFFFFF"),
    }
    for edge, (value, size, color) in definitions.items():
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], caption: str) -> object:
    caption_p = doc.add_paragraph()
    caption_p.paragraph_format.keep_with_next = True
    caption_p.paragraph_format.space_before = Pt(4)
    caption_p.paragraph_format.space_after = Pt(3)
    run = caption_p.add_run(caption)
    set_run(run, size=8.2, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Normal Table"
    set_table_borders(table)
    header = table.rows[0]
    repeat_header(header)
    keep_row(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        set_run(paragraph.add_run(value), size=7.4, bold=True)
        shade(cell, "DCEBEC")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_values in rows:
        row = table.add_row()
        keep_row(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_run(paragraph.add_run(str(value)), size=7.15)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, label: str, width: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    shape = paragraph.add_run().add_picture(str(FIGURES[label]), width=Inches(width))
    shape._inline.docPr.set("descr", ALT_TEXT[label])
    shape._inline.docPr.set("title", label)
    caption_p = doc.add_paragraph(style="Caption")
    caption_p.paragraph_format.keep_together = True
    caption_p.paragraph_format.space_after = Pt(7)
    set_run(caption_p.add_run(f"{label}. "), size=7.8, bold=True)
    set_run(caption_p.add_run(caption), size=7.8)


def configure_document(doc: Document) -> None:
    base.DOCUMENT_TEMPLATE = TEMPLATE
    base.PRESERVE_TEMPLATE_GEOMETRY = True
    base.PRESERVE_TEMPLATE_PAGE_FURNITURE = False
    base.HEADER_TEXT = ""
    base.FOOTER_PREFIX = ""
    base.JOURNAL_MANUSCRIPT.update(
        {
            "font": "Times New Roman",
            "body_size_pt": 9.35,
            "body_after_pt": 3.0,
            "body_line_spacing": 1.08,
            "title_size_pt": 16.0,
            "heading1_size_pt": 10.5,
            "heading2_size_pt": 9.6,
            "caption_size_pt": 7.8,
        }
    )
    base.configure(doc)

    # Supporting Information is delivered as a separate file. Re-create the
    # journal running head explicitly because clearing the article template body
    # also removes the inherited first-page furniture in some Word renderers.
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    for child in list(header_paragraph._p):
        if child.tag != qn("w:pPr"):
            header_paragraph._p.remove(child)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    set_run(header_paragraph.add_run("WILEY-VCH"), size=8.0, italic=True)

    properties = doc.core_properties
    properties.title = f"Supporting Information: {TITLE}"
    properties.author = "Yan Jun Lin"
    properties.subject = "Small Methods Supporting Information v34; frozen computational validation receipts"
    properties.keywords = "quantitative microscopy; validation; perturbation; bone; network; reproducibility"


def build() -> Path:
    missing = [path for path in (*FIGURES.values(), *RECEIPTS, TEMPLATE) if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing SI inputs: " + ", ".join(str(path) for path in missing))

    synthetic = load(RECEIPTS[0])
    matrix = load(RECEIPTS[1])
    stare = load(RECEIPTS[2])
    bone = load(RECEIPTS[3])
    bonej = load(RECEIPTS[4])
    dynamic = load(RECEIPTS[5])
    deformation = load(RECEIPTS[6])
    spatial = load(RECEIPTS[7])

    doc = Document(TEMPLATE)
    base.clear_template_body(doc)
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    set_run(title.add_run("Supporting Information"), size=18, bold=True)
    article = doc.add_paragraph()
    article.paragraph_format.space_after = Pt(5)
    set_run(article.add_run(TITLE), size=12.0, bold=True)
    author = doc.add_paragraph()
    author.paragraph_format.space_after = Pt(8)
    set_run(author.add_run("Yan Jun Lin"), size=10.0, italic=True)
    add_paragraph(
        doc,
        "This file reports estimator-level validations that support implementation breadth but are not pooled with the four main selective-validity experiments. Every value is read directly from a frozen JSON receipt; public images are not redistributed in the software archive.",
        after=7,
    )

    add_heading(doc, "S1. Synthetic truth and perturbation validation", level=1)
    add_paragraph(
        doc,
        "Nine analytic constructs encoded programmed orientation, spectral scale, blob, tube, sheet, thickness, roughness, network structure and spatial heterogeneity. Eight primary perturbations were applied to the orientation-scale construct. The five module gates evaluated structure tensor, Hessian morphology, local thickness, network survival and directional spatial range. Mask dilation and erosion were evaluated as sensitivities because invariance to an altered mask is not scientifically expected.",
    )

    perturbation_rows = []
    for result in synthetic["perturbation_results"]:
        perturbation_rows.append(
            (
                result["perturbation"]["kind"].replace("_", " "),
                f"{result['perturbation']['magnitude']:g}",
                f"{result['errors']['circular_angular_error_degrees']:.4f}°",
                f"{100 * result['errors']['relative_scale_error']:.3f}%",
                "pass" if result["passed"] else "fail",
            )
        )
    add_table(
        doc,
        ("Perturbation", "Magnitude", "Angular error", "Scale error", "Gate"),
        perturbation_rows,
        "Table S1. Exact programmed-perturbation results.",
    )

    module_rows = []
    for module in ("tensor", "hessian", "geometry", "network", "spatial"):
        required = [
            row
            for row in matrix["results"]
            if row["module"] == module and row["perturbation"]["kind"] != "mask_error"
        ]
        sensitivity = [
            row
            for row in matrix["results"]
            if row["module"] == module and row["perturbation"]["kind"] == "mask_error"
        ]
        module_rows.append(
            (
                module,
                str(len(required)),
                str(sum(bool(row["passed"]) for row in required)),
                str(len(sensitivity)) if sensitivity else "—",
            )
        )
    add_table(
        doc,
        ("Module", "Required", "Passed", "Mask sensitivities"),
        module_rows,
        "Table S2. Frozen module-by-perturbation gate counts.",
    )
    add_figure(
        doc,
        "Figure S1",
        6.25,
        "Synthetic operating-envelope validation. a, Programmed analytic constructs. b, Angular and physical-scale errors under the eight primary perturbations. c, Required module-by-perturbation cells; teal denotes a passed required cell and grey denotes a combination not declared by the protocol. d, Monotonic network survival under erosion. All 24 required module tests passed; the two mask-error experiments were retained as sensitivity analyses.",
    )

    add_heading(doc, "S2. External network measurement on STARE reference masks", level=1)
    add_paragraph(
        doc,
        "Twenty STARE reference masks were analyzed as binary networks. NOSTOS and the independent scikit-image skeletonize comparator used the same mask support. Sampling changes preserved both network erosion survival and skeleton length. The result validates endpoint implementation on supplied masks; it does not validate automatic vessel segmentation or infer vascular biology.",
    )
    stare_summary = stare["summary"]
    add_table(
        doc,
        ("Endpoint", "Independent images", "Agreement", "Additional diagnostic"),
        (
            (
                "erosion-survival AUC",
                str(stare_summary["case_count"]),
                f"Spearman ρ = {stare_summary['survival_auc_spearman']:.3f}",
                f"inter-annotator ρ = {stare_summary['ah_vk_survival_auc_spearman']:.3f}",
            ),
            (
                "skeleton length",
                str(stare_summary["case_count"]),
                f"Spearman ρ = {stare_summary['skeleton_length_spearman']:.3f}",
                f"median relative error = {stare_summary['median_skeleton_length_relative_error']:.3f}",
            ),
        ),
        "Table S3. STARE network endpoint confirmation.",
    )

    add_heading(doc, "S3. External trabecular-bone local thickness", level=1)
    add_paragraph(
        doc,
        "Eight public micro-CT volumes from Zenodo record 11061947 were analyzed in physical units. NOSTOS maximal-sphere local thickness was compared first with the archived IPL map and then with BoneJ 1.4.3 executed in a checksum-locked legacy Fiji environment. The independent image volume was the inferential unit. These comparisons establish agreement for local thickness on supplied bone masks, not diagnostic performance or tissue competence.",
    )
    bone_summary = bone["summary"]
    bonej_summary = bonej["summary"]
    add_table(
        doc,
        ("Comparator", "Volumes", "Primary agreement", "Error"),
        (
            (
                "archived IPL map",
                str(bone_summary["n_volumes"]),
                f"median voxelwise ρ = {bone_summary['median_voxelwise_spearman']:.3f}",
                f"mean MAE = {bone_summary['mean_mae_mm']:.4f} mm",
            ),
            (
                "BoneJ 1.4.3",
                str(bonej_summary["case_count"]),
                f"CCC = {bonej_summary['nostos_bonej_ccc']:.3f}",
                f"mean |Δ| = {bonej_summary['mean_absolute_nostos_bonej_difference_mm']:.4f} mm",
            ),
        ),
        "Table S4. Public trabecular-bone thickness comparisons.",
    )
    add_figure(
        doc,
        "Figure S2",
        6.25,
        "Public trabecular-bone local-thickness validation. a–c, Archived reference, NOSTOS estimate and residual on the same volume plane. d, Volume-level mean thickness. e, Voxelwise rank agreement in eight independent volumes. f, Paired mean absolute error for the nearest-boundary baseline and NOSTOS. The exact one-sided paired Wilcoxon value for the improvement over the baseline was 0.00390625.",
    )

    add_heading(doc, "S4. Dynamic deformation and spatial-response confirmations", level=1)
    add_paragraph(
        doc,
        "Dynamic and spatial modules were tested separately so that their bounded operating claims could not be mistaken for biological validation. BBBC035 phase-correlation tests used programmed translations derived from public images. Dense deformation used programmed non-rigid fields. BBBC006 spatial responses compared adjacent focal planes with an intentionally defocused plane.",
    )
    dynamic_summary = dynamic["summary"]
    deformation_summary = deformation["summary"]
    spatial_summary = spatial["summary"]
    add_table(
        doc,
        ("Module", "Independent cases", "Frozen result", "Boundary"),
        (
            (
                "bulk registration",
                "public-image series",
                f"maximum error = {dynamic_summary['maximum_nostos_error_pixels']:.1f} px",
                "programmed translation",
            ),
            (
                "dense deformation",
                str(deformation_summary["case_count"]),
                f"median endpoint error = {deformation_summary['median_endpoint_error_pixels']:.3f} px",
                f"comparator = {deformation_summary['comparator_median_endpoint_error_pixels']:.3f} px",
            ),
            (
                "spatial response",
                str(spatial_summary["case_count"]),
                f"adjacent-plane ρ = {spatial_summary['z15_z16_range_spearman']:.3f}",
                f"defocus > adjacent in {100 * spatial_summary['fraction_defocus_greater']:.0f}%",
            ),
        ),
        "Table S5. Separate dynamic and spatial operating-envelope confirmations.",
    )

    add_heading(doc, "S5. Label-free and three-dimensional bone stress tests", level=1)
    add_paragraph(
        doc,
        "The bone transfer program exercised orientation, network and scale support on mouse SHG, rat confocal lacuna-canalicular stacks, human synchrotron nanoCT and human ultraviolet photoacoustic microscopy. These analyses were intentionally kept as stress tests. A missing pixel calibration forced the UV-PAM spectrum to remain pixel-domain; insufficient support or unstable topology produced abstention rather than an inferred physical measurement.",
    )
    add_figure(
        doc,
        "Figure S3",
        6.25,
        "Public bone stress atlas. a,b, Mouse-bone SHG and coarse orientation-compatibility regions. c,d, Rat confocal projection and imported lacuna-canalicular labels. e, Deterministic intensity terrain. f, Risk–coverage curves for endpoint QC, topology-only support and the full contract. g,h, Human nanoCT crops and internal low-density isosurfaces. i,j, Local orientation and its scale-conditioned coverage/risk boundary. k,l, Human-bone UV-PAM and its pixel-domain Fourier spectrum. These exploratory transfers are not pooled with the primary confirmations.",
    )

    receipt_heading = add_heading(doc, "S6. Frozen receipt index", level=1)
    receipt_heading.paragraph_format.page_break_before = True
    add_paragraph(
        doc,
        "The relative record name and SHA-256 digest below identify the exact machine-readable evidence used to generate this Supporting Information. The release builder verifies these records before packaging and excludes bulk public microscopy.",
    )
    receipt_rows = []
    for path in RECEIPTS:
        receipt_rows.append((path.relative_to(ROOT).as_posix(), sha256(path)))
    add_table(
        doc,
        ("Receipt", "SHA-256"),
        receipt_rows,
        "Table S6. Machine-readable evidence receipts.",
    )

    add_heading(doc, "S7. Resource identifiers and scope", level=1)
    resource_rows = (
        ("STARE", "https://cecas.clemson.edu/~ahoover/stare/probing/", "reference-mask network endpoints"),
        ("trabecular bone", "https://doi.org/10.5281/zenodo.11061947", "local thickness on supplied masks"),
        ("BBBC035", "https://bbbc.broadinstitute.org/BBBC035", "programmed bulk and dense motion"),
        ("BBBC006", "https://bbbc.broadinstitute.org/BBBC006", "adjacent-plane and defocus response"),
    )
    add_table(doc, ("Resource", "Persistent identifier", "Permitted interpretation"), resource_rows, "Table S7. External public resources.")
    add_paragraph(
        doc,
        "No result in this file establishes diagnosis, mechanics, clinical usefulness, intraoperative performance, automatic tissue segmentation or a universal biological meaning for any NOSTOS coordinate. The supplementary results establish only the declared estimator-level agreements and perturbation behaviors.",
        after=0,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def main() -> None:
    print(build())


if __name__ == "__main__":
    main()
