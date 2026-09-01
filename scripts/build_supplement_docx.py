"""Build NOSTOS supplementary information with secondary analyses."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from build_manuscript_docx import ROOT, PUB, OUTPUT, configure, font, caption, add_table

OUT=ROOT/'docs'/'NOSTOS_supplementary_information.docx'

def picture(doc,path,width=6.35,alt="Supplementary analysis figure"):
 p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.keep_with_next=True;p.add_run().add_picture(str(path),width=Inches(width));drawing=doc.inline_shapes[-1]._inline.docPr;drawing.set("descr",alt);drawing.set("title",path.stem.replace('_',' '))

def build():
 doc=Document();configure(doc)
 p=doc.add_paragraph();font(p.add_run('NOSTOS Supplementary Information'),18,bold=True)
 p=doc.add_paragraph('Secondary analyses, robustness experiments, and complete numerical results');font(p.runs[0],10,italic=True)
 doc.add_heading('Supplementary Figures',level=1)
 picture(doc,PUB/'figure_1.jpg',alt='Medial-site angular-entropy associations with histologic scores.');caption(doc,'Supplementary Figure','S1a','Medial-site entropy associations retained as a conventional scatterplot view.')
 picture(doc,PUB/'figure_2.jpg',alt='Lateral-site angular-entropy associations with histologic scores.');caption(doc,'Supplementary Figure','S1b','Lateral-site entropy associations retained as a conventional scatterplot view.')
 picture(doc,PUB/'figure_3.jpg',4.5,alt='Observed and out-of-fold predicted PLM scores using angular entropy.');caption(doc,'Supplementary Figure','S2','Participant-grouped nested-cross-validation predictions for PLM using angular entropy alone.')
 picture(doc,PUB/'figure_4.jpg',alt='Acquisition perturbation and cartilage-mask boundary sensitivity results.');caption(doc,'Supplementary Figure','S3','Acquisition perturbation and cartilage-mask boundary sensitivity analyses.')
 doc.add_heading('Supplementary Tables',level=1)
 for n,label in ((3,'S3'),(4,'S4'),(5,'S5')): add_table(doc,n,label)
 doc.core_properties.title='NOSTOS Supplementary Information';doc.core_properties.author='NOSTOS study team';doc.save(OUT);print(OUT)
if __name__=='__main__':build()
