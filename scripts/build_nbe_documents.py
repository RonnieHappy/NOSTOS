"""Build the evidence-gated NOSTOS Nature Biomedical Engineering package."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = ROOT / "outputs" / "nbe_package"
FIG = OUT / "figures"
FONT = "Times New Roman"
INK = RGBColor(25, 31, 35)
MUTED = RGBColor(89, 98, 103)
BLUE = RGBColor(27, 75, 96)
RED = RGBColor(155, 28, 28)
HEADER_FILL = "E8EEF1"
CAUTION_FILL = "FFF4D6"
RISK_FILL = "FCE8E6"
WIDTH_DXA = 9360


def set_font(run, size=10.5, *, bold=None, italic=None, color=INK):
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)


def table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    tblw.set(qn("w:w"), str(WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tblpr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcw = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)


def configure(doc, running_label):
    section = doc.sections[0]
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16
    for name, size, before, after, color in (
        ("Heading 1", 15, 16, 7, BLUE),
        ("Heading 2", 12, 12, 5, BLUE),
        ("Heading 3", 10.5, 9, 4, INK),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(running_label), 8, italic=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def inline(paragraph, text, size=10.5):
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**"):
            set_font(paragraph.add_run(part[2:-2]), size, bold=True)
        elif part.startswith("*"):
            set_font(paragraph.add_run(part[1:-1]), size, italic=True)
        elif part.startswith("`"):
            set_font(paragraph.add_run(part[1:-1]), size - .5, color=BLUE)
        else:
            set_font(paragraph.add_run(part), size)


def add_status_box(doc, text, risk=False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_width(table, [WIDTH_DXA])
    cell = table.cell(0, 0)
    shade(cell, RISK_FILL if risk else CAUTION_FILL)
    cell_margins(cell, 110, 110, 140, 140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    inline(paragraph, text, 9.5)
    for run in paragraph.runs:
        run.font.color.rgb = RED if risk else INK


def add_markdown_table(doc, rows):
    headers = rows[0]
    data = rows[2:] if len(rows) > 1 and set(rows[1]) <= {"---", ":---", "---:", ":---:"} else rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [WIDTH_DXA // len(headers)] * len(headers)
    widths[-1] += WIDTH_DXA - sum(widths)
    for c, value in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = value
        shade(cell, HEADER_FILL)
    for values in data:
        cells = table.add_row().cells
        for c, value in enumerate(values):
            cells[c].text = value
    table_width(table, widths)
    for ridx, row in enumerate(table.rows):
        for cidx, cell in enumerate(row.cells):
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx in (0, len(row.cells) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, 8.2, bold=ridx == 0)


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(6.5))
    drawing = doc.inline_shapes[-1]._inline.docPr
    drawing.set("descr", caption)
    drawing.set("title", path.stem.replace("_", " "))
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(8)
    inline(cp, caption, 8.7)


def build_from_markdown(source, output, running_label):
    doc = Document()
    configure(doc, running_label)
    lines = source.read_text(encoding="utf-8").splitlines()
    table_rows = []
    for line in lines + [""]:
        if line.startswith("|") and line.endswith("|"):
            table_rows.append([part.strip() for part in line.strip("|").split("|")])
            continue
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            set_font(p.add_run(line[2:]), 18, bold=True, color=INK)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("[DATA REQUIRED"):
            add_status_box(doc, line, risk=True)
        elif line == "[[FIGURE:roadmap]]":
            add_picture(doc, FIG / "figure_1_evidence_gates.png", "Figure 1 | Evidence-gated NOSTOS translation. Green denotes completed public-histology evidence; red denotes experiments required before an NBE submission.")
        elif line == "[[FIGURE:public]]":
            add_picture(doc, ROOT / "outputs" / "main_figures" / "figure_1_mega.jpg", "Figure 2 | Completed public-histology foundation. Repository-derived human microscopy, algorithmic tissue proposal, spatial-frequency mapping and participant-level associations.")
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            inline(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            inline(p, line)
    props = doc.core_properties
    props.title = lines[0].removeprefix("# ")
    props.author = "NOSTOS study team"
    props.subject = "Evidence-gated translational manuscript development package"
    props.comments = "Future results are explicitly marked DATA REQUIRED—DO NOT SUBMIT."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_roadmap():
    FIG.mkdir(parents=True, exist_ok=True)
    mpl.rcParams.update({"font.family": "serif", "font.serif": [FONT], "font.size": 8, "savefig.dpi": 600})
    fig, ax = plt.subplots(figsize=(7.2, 3.3))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    stages = [
        ("Public histology", "90 participants\nserial-section replication", "#D8ECE6", "COMPLETE"),
        ("Measurement", "PS-OCT/OCE\nrepeatability", "#FBE6E3", "REQUIRED"),
        ("Mechanics", "registered local\nindentation", "#FBE6E3", "REQUIRED"),
        ("Specificity", "matrix-selective\nperturbations", "#FBE6E3", "REQUIRED"),
        ("Translation", "cadaver + prospective\nboundary validation", "#FBE6E3", "REQUIRED"),
    ]
    for index, (title, body, fill, status) in enumerate(stages):
        x = .01 + index * .2
        box = FancyBboxPatch((x, .37), .17, .42, boxstyle="round,pad=0.008,rounding_size=0.012", facecolor=fill, edgecolor="#7E8B90", linewidth=.8)
        ax.add_patch(box)
        ax.text(x + .085, .68, title, ha="center", va="center", fontweight="bold", fontsize=8.4)
        ax.text(x + .085, .53, body, ha="center", va="center", fontsize=7.2)
        ax.text(x + .085, .405, status, ha="center", va="center", fontsize=6.5, color="#246B55" if status == "COMPLETE" else "#9B1C1C", fontweight="bold")
        if index < 4:
            ax.add_patch(FancyArrowPatch((x + .174, .58), (x + .198, .58), arrowstyle="-|>", mutation_scale=9, color="#315E70", linewidth=1))
    ax.text(.01, .93, "NOSTOS: evidence must cross every gate before clinical claims", fontsize=12, fontweight="bold", color="#193947")
    ax.text(.01, .15, "Primary translational estimand", fontsize=8, fontweight="bold", color="#193947")
    ax.text(.01, .07, "Within-specimen association between a frozen architecture score and registered local equilibrium modulus", fontsize=8)
    fig.savefig(FIG / "figure_1_evidence_gates.png", bbox_inches="tight", facecolor="white")
    fig.savefig(FIG / "figure_1_evidence_gates.svg", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_roadmap()
    build_from_markdown(DOCS / "nbe_flagship_manuscript.md", DOCS / "NOSTOS_NBE_flagship_living_article.docx", "NOSTOS | NBE flagship development Article")
    build_from_markdown(DOCS / "nbe_supplementary_methods.md", DOCS / "NOSTOS_NBE_supplementary_methods.docx", "NOSTOS | Supplementary Methods")
    build_from_markdown(DOCS / "nbe_presubmission_synopsis.md", DOCS / "NOSTOS_NBE_presubmission_synopsis.docx", "NOSTOS | Presubmission development synopsis")
    build_from_markdown(DOCS / "nbe_claim_evidence_ledger.md", DOCS / "NOSTOS_NBE_claim_evidence_ledger.docx", "NOSTOS | Claim-evidence ledger")
    shutil.copy2(DOCS / "NOSTOS_manuscript_embedded.docx", DOCS / "NOSTOS_public_histology_article.docx")
    shutil.copy2(DOCS / "NOSTOS_supplementary_information.docx", DOCS / "NOSTOS_public_histology_supplement.docx")
    print("Built NOSTOS two-track manuscript package")


if __name__ == "__main__":
    main()
