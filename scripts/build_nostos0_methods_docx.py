"""Build the evidence-locked NOSTOS-0 computational-methods Article."""
from __future__ import annotations

import re
import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "NOSTOS0_SOFTWARE_RESOURCE_ARTICLE.md"
OUTPUT = ROOT / "docs" / "NOSTOS0_computational_methods_submission_candidate_v30.docx"
DOCUMENT_TEMPLATE: Path | None = None
PRESERVE_TEMPLATE_GEOMETRY = False
PRESERVE_TEMPLATE_PAGE_FURNITURE = False
FIGURE_CAPTION_SEPARATOR = " | "
FIGURES = {
    "Figure 1 near here": ROOT / "figures/nostos0/figure_1_validity_contract_atlas.png",
    "Figure 2 near here": ROOT / "figures/nostos0/figure_2_biosr_selective_validity.png",
    "Figure 3 near here": ROOT / "figures/nostos0/figure_3_fmd_hidden_failure.png",
    "Figure 4 near here": ROOT / "figures/nostos0/figure_4_fmd_hierarchical_confirmation.png",
    "Supplementary Figure 1 near here": ROOT / "figures/nostos0/supplementary_figure_1_bone_contract_stress.png",
}
FIGURE_ALT = {
    "Figure 1": "NOSTOS validity-contract atlas showing public BioSR and FMD microscopy, deterministic structural maps, the compile-to-decision workflow, a scale response, an input-only validity fingerprint and a supported measurement lattice.",
    "Figure 2": "BioSR selective-validity result showing paired public microscopy, degradation outcomes, field-paired invalid-emission risk, risk-coverage curves and enrichment of invalid values among NOSTOS-only rejections.",
    "Figure 3": "FMD failure-localization result showing the capture-averaging image ladder, conditional risk surface, development and confirmation failure matrices and the failure-preserving repair sequence.",
    "Figure 4": "FMD hierarchical-confirmation result showing four untouched fields, frozen support cells, prospective emissions, matched acquisition-QC errors, risk-coverage behavior, bootstrap uncertainty and exact zero-event bounds.",
    "Supplementary Figure 1": "Public unstained and label-free bone imaging stress tests across two-dimensional and three-dimensional acquisition families.",
}
FIGURE_WIDTHS = {
    "Figure 1": 6.65,
    "Figure 2": 5.75,
    "Figure 3": 5.75,
    "Figure 4": 5.75,
    "Supplementary Figure 1": 6.35,
}
HEADER_TEXT = "NOSTOS-0 | computational methods Article"
HEADER_ALIGNMENT = WD_ALIGN_PARAGRAPH.RIGHT
FOOTER_PREFIX = "Page "
DOC_TITLE = "NOSTOS compiles acquisition- and scale-specific validity contracts for quantitative microscopy"
DOC_SUBJECT = "Evidence-locked NOSTOS-0 computational-methods submission candidate v30"
DOC_AUTHOR = "Yany Lin"
DOC_KEYWORDS = "computational microscopy; measurement validity; selective prediction; abstention; reproducibility"

# The narrative_proposal preset is resolved here with a named
# journal_manuscript override required by the author: Times New Roman throughout,
# compact journal spacing and wider printable image area.
JOURNAL_MANUSCRIPT = {
    "page_width_in": 8.5,
    "page_height_in": 11.0,
    "margins_in": {"top": 0.72, "right": 0.78, "bottom": 0.72, "left": 0.78},
    "header_distance_in": 0.30,
    "footer_distance_in": 0.30,
    "font": "Times New Roman",
    "body_size_pt": 10.0,
    "body_after_pt": 5.0,
    "body_line_spacing": 1.06,
    "title_size_pt": 18.0,
    "heading1_size_pt": 13.0,
    "heading2_size_pt": 11.0,
    "caption_size_pt": 8.5,
}


def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = RGBColor(*color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def force_style_font(style, family: str = "Times New Roman") -> None:
    """Set all script fonts and remove theme overrides that LibreOffice may prefer."""
    style.font.name = family
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attribute in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        qname = qn(attribute)
        if qname in fonts.attrib:
            del fonts.attrib[qname]
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), family)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    if not PRESERVE_TEMPLATE_GEOMETRY:
        section.page_width = Inches(JOURNAL_MANUSCRIPT["page_width_in"])
        section.page_height = Inches(JOURNAL_MANUSCRIPT["page_height_in"])
        section.top_margin = Inches(JOURNAL_MANUSCRIPT["margins_in"]["top"])
        section.bottom_margin = Inches(JOURNAL_MANUSCRIPT["margins_in"]["bottom"])
        section.left_margin = Inches(JOURNAL_MANUSCRIPT["margins_in"]["left"])
        section.right_margin = Inches(JOURNAL_MANUSCRIPT["margins_in"]["right"])
        section.header_distance = Inches(JOURNAL_MANUSCRIPT["header_distance_in"])
        section.footer_distance = Inches(JOURNAL_MANUSCRIPT["footer_distance_in"])
    def paragraph_style(name: str):
        try:
            return doc.styles[name]
        except KeyError:
            return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    normal = doc.styles["Normal"]
    force_style_font(normal, JOURNAL_MANUSCRIPT["font"])
    normal.font.size = Pt(JOURNAL_MANUSCRIPT["body_size_pt"])
    normal.paragraph_format.space_after = Pt(JOURNAL_MANUSCRIPT["body_after_pt"])
    normal.paragraph_format.line_spacing = JOURNAL_MANUSCRIPT["body_line_spacing"]
    styles_element = doc.styles.element
    defaults = styles_element.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, defaults)
    run_defaults = defaults.find(qn("w:rPrDefault"))
    if run_defaults is None:
        run_defaults = OxmlElement("w:rPrDefault")
        defaults.insert(0, run_defaults)
    default_rpr = run_defaults.find(qn("w:rPr"))
    if default_rpr is None:
        default_rpr = OxmlElement("w:rPr")
        run_defaults.append(default_rpr)
    default_fonts = default_rpr.find(qn("w:rFonts"))
    if default_fonts is None:
        default_fonts = OxmlElement("w:rFonts")
        default_rpr.insert(0, default_fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        default_fonts.set(qn(attribute), "Times New Roman")
    for name, size, before, after in (
        ("Title", JOURNAL_MANUSCRIPT["title_size_pt"], 0, 8),
        ("Heading 1", JOURNAL_MANUSCRIPT["heading1_size_pt"], 11, 4),
        ("Heading 2", JOURNAL_MANUSCRIPT["heading2_size_pt"], 8, 3),
    ):
        style = paragraph_style(name)
        force_style_font(style); style.font.size = Pt(size)
        style.font.bold = True; style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = paragraph_style("Caption")
    force_style_font(caption); caption.font.size = Pt(JOURNAL_MANUSCRIPT["caption_size_pt"])
    caption.font.bold = False; caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.space_before = Pt(3); caption.paragraph_format.space_after = Pt(7)
    force_style_font(doc.styles["Header"])
    force_style_font(doc.styles["Footer"])
    if not PRESERVE_TEMPLATE_PAGE_FURNITURE:
        header = section.header.paragraphs[0]
        for child in list(header._p):
            if child.tag != qn("w:pPr"):
                header._p.remove(child)
        header.alignment = HEADER_ALIGNMENT
        if HEADER_TEXT:
            set_font(header.add_run(HEADER_TEXT), 8, italic=True, color=(95, 95, 95))
        footer = section.footer.paragraphs[0]
        for child in list(footer._p):
            if child.tag != qn("w:pPr"):
                footer._p.remove(child)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if FOOTER_PREFIX:
            set_font(footer.add_run(FOOTER_PREFIX), 8, color=(95, 95, 95))
        add_page_field(footer)


def clear_template_body(doc: Document) -> None:
    """Remove template placeholders while preserving styles, relationships and sectPr."""
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_inline(paragraph, text: str) -> None:
    text = (text.replace("\\(", "").replace("\\)", "")
            .replace("\\rho", "ρ").replace("\\Delta", "Δ")
            .replace("\\theta", "θ").replace("\\tau", "τ")
            .replace("\\ell", "ℓ").replace("\\Phi", "Φ")
            .replace("\\mu", "µ").replace("\\", ""))
    token = re.compile(
        r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[(?:\d+(?:[–-]\d+)?(?:,\s*\d+(?:[–-]\d+)?)*)\])"
    )
    for part in token.split(text):
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), bold=True)
        elif part.startswith("*") and part.endswith("*"):
            set_font(paragraph.add_run(part[1:-1]), italic=True)
        elif part.startswith("`") and part.endswith("`"):
            set_font(paragraph.add_run(part[1:-1]), 9, color=(55, 55, 55))
        elif part.startswith("[") and part.endswith("]"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, 8)
            run.font.superscript = True
        else:
            set_font(paragraph.add_run(part))


def logical_lines(text: str) -> list[str]:
    """Collapse hard-wrapped Markdown prose into semantic paragraphs.

    The manuscript source is intentionally wrapped for review in a text editor.
    Treating every physical line as a Word paragraph produces ragged half-width
    text and large blank regions. Blank lines and Markdown control lines define
    paragraph boundaries; numbered references are kept as separate entries.
    """
    output: list[str] = []
    buffer: list[str] = []
    in_references = False

    def flush() -> None:
        if buffer:
            output.append(" ".join(buffer))
            buffer.clear()

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue

        is_control = (
            line.startswith("# ")
            or line.startswith("## ")
            or line.startswith("### ")
            or line in {"---", r"\[", r"\]"}
            or line.startswith("**[")
        )
        if is_control:
            flush()
            output.append(line)
            in_references = line == "## References"
            continue

        if in_references and re.match(r"^\d+\.\s", line):
            flush()
            buffer.append(line)
            continue

        buffer.append(line)

    flush()
    return output


def build(source: Path = SOURCE, output: Path = OUTPUT) -> Path:
    lines = logical_lines(source.read_text(encoding="utf-8"))
    legends: dict[str, str] = {}
    for line in lines:
        match = re.match(r"\*\*((?:Supplementary )?Figure \d+)(?: \||\.) (.+?)\.\*\*\s*(.*)", line)
        if match:
            legends[match.group(1)] = (
                f"{match.group(1)}{FIGURE_CAPTION_SEPARATOR}{match.group(2)}. {match.group(3)}"
            )

    doc = Document(DOCUMENT_TEMPLATE) if DOCUMENT_TEMPLATE is not None else Document()
    if DOCUMENT_TEMPLATE is not None:
        clear_template_body(doc)
    configure(doc)
    in_legend_section = False; in_references = False; in_equation = False; equation_lines = []
    for raw in lines:
        line = raw.strip()
        if line == "## Figure legends":
            in_legend_section = True
            continue
        if in_legend_section:
            if line == "## References":
                in_legend_section = False
                doc.add_heading("References", level=1)
                in_references = True
                continue
            else:
                continue
        if not line: continue
        if line == "\\[": in_equation = True; equation_lines = []; continue
        if line == "\\]":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)
            set_font(p.add_run("Φ(I, M, Δ) = {Φₘ(ℓ, θ, τ, r)}ₘ₌₁ᴹ"), 10, italic=True)
            set_font(p.add_run("\t(1)"), 9)
            in_equation = False; continue
        if in_equation: equation_lines.append(line); continue
        if line.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            set_font(p.add_run(line[2:]), JOURNAL_MANUSCRIPT["title_size_pt"], bold=True)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2); continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1); continue
        figure_key = next((key for key in sorted(FIGURES, key=len, reverse=True) if key in line), None)
        if figure_key:
            label = figure_key.removesuffix(" near here")
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True; p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
            figure_width = FIGURE_WIDTHS.get(label, 6.35)
            shape = p.add_run().add_picture(str(FIGURES[figure_key]), width=Inches(figure_width))
            shape._inline.docPr.set("descr", FIGURE_ALT[label])
            shape._inline.docPr.set("title", label)
            if label in legends:
                cap = doc.add_paragraph(style="Caption"); cap.paragraph_format.keep_together = True
                add_inline(cap, legends[label])
            continue
        if line == "---": continue
        p = doc.add_paragraph()
        if line.startswith("All biological images remain in their originating public repositories."):
            p.paragraph_format.keep_together = True
        if line.startswith("**Article type:**") or line.startswith("**Target:**") or line.startswith("**Version:**") or line.startswith("**Status:**"):
            p.paragraph_format.space_after = Pt(1)
        if re.match(r"^\d+\. ", line):
            p.paragraph_format.left_indent = Inches(.18)
            p.paragraph_format.first_line_indent = Inches(-.18)
            p.paragraph_format.space_after = Pt(1 if in_references else 2)
            if in_references:
                p.paragraph_format.line_spacing = 1.0
        add_inline(p, line)
        if in_references:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_font(run, 8.5)

    props = doc.core_properties
    props.title = DOC_TITLE
    props.subject = DOC_SUBJECT
    props.author = DOC_AUTHOR
    props.keywords = DOC_KEYWORDS
    doc.save(output)
    return output


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    print(build(args.source, args.output))
