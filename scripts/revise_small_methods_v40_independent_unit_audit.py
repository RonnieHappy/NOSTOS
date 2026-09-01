"""Create the v40 manuscript by applying the independent-unit risk audit to v39."""

from __future__ import annotations

import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscripts/NOSTOS_Small_Methods_submission_ready_v39.docx"
OUTPUT = ROOT / "manuscripts/NOSTOS_Small_Methods_submission_ready_v40.docx"


def set_tnr(run) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")


def replace_single_run(paragraph, expected: str, replacement: str) -> None:
    if paragraph.text != expected or len(paragraph.runs) != 1:
        raise RuntimeError(f"Unexpected source paragraph: {paragraph.text[:120]!r}")
    paragraph.runs[0].text = replacement
    set_tnr(paragraph.runs[0])


def main() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    p41 = document.paragraphs[41]
    replace_single_run(
        p41,
        p41.text,
        p41.text
        + " A retrospective independent-unit audit of the same frozen policy found at least one accepted invalid condition in 7 of 24 regions (29.2%); the one-sided 95% exact Clopper-Pearson upper bound was 47.9%. The condition-level reduction therefore does not establish a 20% region-level risk guarantee.",
    )

    p44 = document.paragraphs[44]
    caption_addition = (
        " Seven of 24 independent regions contained at least one accepted invalid condition; "
        "the one-sided 95% exact upper bound was 47.9%, so these nested-condition results are descriptive rather than a finite-sample region-level guarantee."
    )
    if caption_addition.strip() in p44.text:
        raise RuntimeError("Figure 5 audit sentence is already present.")
    run = p44.add_run(caption_addition)
    set_tnr(run)

    p58 = document.paragraphs[58]
    replacement_58 = p58.text.replace(
        "The PSHG-TISS result uses programmed shifts in one microscope family.",
        "The PSHG-TISS result uses programmed shifts in one microscope family. Although the frozen policy retained 7 invalid condition rows among 230 accepted rows, 7 of 24 independent regions contained at least one accepted invalid condition and the one-sided 95% exact upper bound was 47.9%; this blocks a finite-sample region-level risk-control claim.",
    )
    if replacement_58 == p58.text:
        raise RuntimeError("PSHG limitations anchor was not found.")
    replace_single_run(p58, p58.text, replacement_58)

    p111 = document.paragraphs[111]
    replacement_111 = (
        "All microscopy data remain in their originating public repositories. FMD is available under CC BY-SA 4.0 at DOI 10.7274/r0-ed2r-4052; BioSR is available at DOI 10.6084/m9.figshare.13264793; PSHG-TISS is available at DOI 10.17605/OSF.IO/UDTQP; the tendon pSHG-XRD resource is available under CC BY 4.0 at DOI 10.5281/zenodo.10979115. Dataset identifiers, licences, archive and member hashes, frozen selection rules and exact commands are included in the software evidence record. Source code is publicly available at https://github.com/RonnieHappy/NOSTOS under the BSD 3-Clause License. The exact review archive contains 1,072 files, has SHA-256 bcefad6b9d94b1da3f084d10cce171bdb8c0ac750eaa862b82046944e3a07942, and passed clean-room verification with 427 tests passed, 8 optional dependency skips and 0 failures. A versioned archival DOI remains required before publication."
    )
    replace_single_run(p111, p111.text, replacement_111)

    fixed_time = datetime(2026, 9, 1, 5, 0, 0, tzinfo=timezone.utc)
    document.core_properties.modified = fixed_time
    document.core_properties.last_modified_by = "NOSTOS reproducibility audit"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
