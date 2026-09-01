"""Build the submission-style NOSTOS manuscript with inline tables and figures."""
from __future__ import annotations

import csv, re
from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "manuscript_draft.md"
PUB = ROOT / "outputs" / "cpu_pilot" / "publication"
MAIN = ROOT / "outputs" / "main_figures"
OUTPUT = ROOT / "docs" / "NOSTOS_public_histology_article_NBE_visual_refinement.docx"
FONT, INK, MUTED, HEADER = "Times New Roman", RGBColor(24,24,24), RGBColor(90,90,90), "E9EFF1"

TABLE_CAPTIONS = {
  1: "Whole-section processing yield and eligible-tile counts by stain and site.",
  2: "Participant-level associations between Safranin-O angular entropy and expert histologic outcomes. Confidence intervals are participant-bootstrap intervals; q values use Benjamini-Hochberg correction within site.",
  3: "Paired-site partial-rank associations adjusted for age, sex, surgical side, cartilage fraction, bone fraction, and analyzed tile count.",
  4: "Participant-grouped nested cross-validation ablation results in 87 paired participants.",
  5: "Frozen component-level associations across first and untouched second sections. Asterisks denote global FDR q<0.05.",
}
FIGURE_CAPTIONS = {
  1: "Microscopy-to-phenotype analysis. (a) Whole human osteochondral Safranin-O section. (b) Conservative cartilage proposal overlay. (c) Binary cartilage proposal and (d) its boundary projected on the source image. (e,g,i) Real cartilage tiles spanning the observed angular-entropy distribution and (f,h,j) their paired Fourier-power spectra. (k) Spatial map of measured cartilage-tile entropy. (l) The same field rendered as a feature-height terrain; height represents angular entropy rather than physical tissue topography. (m) Participant-level associations with PLM, OARSI and HHGS outcomes.",
  2: "Multiscale spectral anatomy. (a–c) Three real cartilage fields spanning the observed entropy distribution, each shown at four progressively reduced sampling scales and paired with its measured log Fourier-power spectrum. (d) Normalized axial spectral profiles across scale. Pixel-grid structure at coarse scales is a resampling consequence, displayed to expose rather than conceal acquisition sensitivity.",
  3: "Orientation cartography. (a) Local dominant-orientation glyphs projected on the specimen. (b) Anisotropy-weighted axial orientation distribution. (c) Entropy-ordered spiral map, with marker area encoding anisotropy. (d) Interpolated orientation streamlines colored by anisotropy; these are a visualization of the estimated field, not physical flow. (e) Spatial orientation glyph field.",
  4: "Spatial topology of the entropy field. (a) Entropy overlaid on the source specimen, (b) iso-entropy contours, and (c) the same values rendered as a feature-height surface. (d) Delaunay neighborhood graph colored by measured tile entropy. (e) Threshold-persistence barcode across sampled field locations. (f) Number of connected high-entropy components across thresholds. Height and connectivity are computational representations, not physical tissue topography.",
  5: "Matched-modality biological context. (a) Safranin-O bright-field and (b) polarized-light microscopy from the same participant. (c) Four relative-depth crop pairs, arranged Safranin-O then PLM, for visual comparison but not deformably registered. (d) Component-association matrix across sites and serial sections. (e) Medial second-section estimates and bootstrap intervals.",
  6: "Serial-section replication. (a,b) Untouched adjacent Safranin-O sections. (c) Common-size alpha composite for visual comparison only; no deformable registration was performed. (d,e) Spatial entropy maps. (f) Section-to-section agreement, (g) difference-versus-mean structure, and (h) selected participant-level paired trajectories.",
  7: "Perturbation and falsification phase space. (a–f) The same measured tile under blur, downsampling, rotation and noise perturbations. (g) Cartilage-mask boundary contour family. (h) Row-normalized 95th-percentile drift matrix across five feature families and the stored acquisition perturbations. (i) Gaussian density implied by the stored permutation-null mean and standard deviation, with observed error marked in red; this is a parametric summary, not an empirical permutation histogram.",
  8: "Participant phenospace and prediction. (a) Principal-component projection of standardized image features, colored by entropy. (b) Kernel-density representation of that phenospace. (c) Medial-to-lateral participant trajectories. (d) Out-of-fold binary calibration and (e) corresponding probability distributions by outcome.",
}

def font(run, size=10.5, bold=None, italic=None, color=INK):
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"): fonts.set(qn(f"w:{key}"), FONT)
    run.font.size = Pt(size); run.font.color.rgb = color
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def inline(paragraph, text, size=10.5):
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^]]+\]\([^)]+\))", text):
        if not part: continue
        link = re.fullmatch(r"\[([^]]+)\]\(([^)]+)\)", part)
        if link: font(paragraph.add_run(f"{link.group(1)} ({link.group(2)})"), size, color=RGBColor(35,80,92))
        elif part.startswith("**"): font(paragraph.add_run(part[2:-2]), size, bold=True)
        elif part.startswith("*"): font(paragraph.add_run(part[1:-1]), size, italic=True)
        elif part.startswith("`"): font(paragraph.add_run(part[1:-1]), size=9)
        else: font(paragraph.add_run(part), size)

def shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(shd)

def margins(cell, top=70, bottom=70, start=90, end=90):
    tcpr = cell._tc.get_or_add_tcPr(); mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None: mar = OxmlElement("w:tcMar"); tcpr.append(mar)
    for side, value in (("top",top),("bottom",bottom),("start",start),("end",end)):
        node=OxmlElement(f"w:{side}"); node.set(qn("w:w"),str(value)); node.set(qn("w:type"),"dxa"); mar.append(node)

def repeat_header(row):
    trpr=row._tr.get_or_add_trPr(); node=OxmlElement("w:tblHeader"); node.set(qn("w:val"),"true"); trpr.append(node)

def set_cell_width(cell, width):
    tcpr=cell._tc.get_or_add_tcPr(); tcw=tcpr.first_child_found_in("w:tcW"); tcw.set(qn("w:w"),str(width)); tcw.set(qn("w:type"),"dxa")

def configure(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(.78); sec.bottom_margin=Inches(.75); sec.left_margin=sec.right_margin=Inches(.82); sec.header_distance=Inches(.32); sec.footer_distance=Inches(.35)
    normal=doc.styles["Normal"]; normal.font.name=FONT; normal.font.size=Pt(10.5); normal.font.color.rgb=INK; normal.paragraph_format.space_after=Pt(5.5); normal.paragraph_format.line_spacing=1.12
    for name,size,before,after in (("Heading 1",14,15,6),("Heading 2",11.5,11,4),("Heading 3",10.5,8,3)):
        s=doc.styles[name]; s.font.name=FONT; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=INK; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run("NOSTOS | Research Article"),8,italic=True,color=MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)

def caption(doc, kind, number, text, above=False):
    p=doc.add_paragraph(); p.paragraph_format.keep_with_next=above; p.paragraph_format.space_before=Pt(6 if above else 3); p.paragraph_format.space_after=Pt(3 if above else 8)
    font(p.add_run(f"{kind} {number} | "),9,bold=True); inline(p,text,9)

def table_rows(number):
    names={1:"table_1_processing_yield.csv",2:"table_2_entropy_associations.csv",3:"table_3_adjusted_associations.csv",4:"table_4_nested_cv_ablations.csv",5:"table_5_mechanistic_associations.csv"}
    rows=list(csv.reader((PUB/names[number]).open(encoding="utf-8")))
    if number==3:
        keep=[i for i,v in enumerate(rows[0]) if v != "covariates"]; rows=[[r[i] for i in keep] for r in rows]
    if number==5:
        hdr=rows[0]; idx={v:i for i,v in enumerate(hdr)}; keep=["site","section_rank","component","n","spearman_rho","bootstrap_ci_lower","bootstrap_ci_upper","q_value_bh_global"]
        data=[[r[idx[k]] for k in keep] for r in rows[1:] if r[idx["feature"]]=="angular_entropy_median"]
        rows=[keep]+data
    return rows

def add_table(doc, number, display_number=None):
    caption(doc,"Table",display_number or number,TABLE_CAPTIONS[number],above=True); rows=table_rows(number); cols=len(rows[0])
    table=doc.add_table(rows=1,cols=cols); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False; table.style="Table Grid"
    widths={1:[2600,1500,1500,2600],2:[900,950,600,1250,1250,1250,1050,1050],3:[1250,1300,650,1150,1150,1150,1150,1450],4:[1050,1500,600,850,850,850,850,850,850,920],5:[720,650,2050,500,900,1000,1000,900]}[number]
    widths=[round(w*9360/sum(widths)) for w in widths]; widths[-1]+=9360-sum(widths)
    for row_index,row in enumerate(rows):
        cells=table.rows[0].cells if row_index==0 else table.add_row().cells
        for i,value in enumerate(row):
            if row_index:
                try: value=f"{float(value):.3g}"
                except ValueError: value=value.replace("_"," ")
                if number==5 and i==4 and float(row[7]) < .05: value += "*"
            else: value=value.replace("bootstrap_ci_lower","CI lower").replace("bootstrap_ci_upper","CI upper").replace("q_value_bh_global","global q").replace("spearman_rho","rho").replace("_"," ")
            cells[i].text=value; set_cell_width(cells[i],widths[i]); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if i in (0,2) else WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
            for run in p.runs: font(run,7.2 if number==5 else 7.5,bold=row_index==0)
            if row_index==0: shade(cells[i],HEADER)
        if row_index==0: repeat_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_figure(doc, number):
    p=doc.add_paragraph(); p.paragraph_format.keep_with_next=True; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(0)
    p.add_run().add_picture(str(MAIN/f"figure_{number}_mega.jpg"),width=Inches(6.5))
    drawing=doc.inline_shapes[-1]._inline.docPr;drawing.set("descr",FIGURE_CAPTIONS[number]);drawing.set("title",f"NOSTOS Figure {number}")
    caption(doc,"Figure",number,FIGURE_CAPTIONS[number])

def build():
    doc=Document(); configure(doc); lines=SOURCE.read_text(encoding="utf-8").splitlines(); first=True; skip_status=False; skip_legends=False; references=False; section=""
    retained_subheads={
      "Methods": {
        "Data source and integrity audit":"Study design and image data",
        "CPU tissue-region proposal":"Image analysis",
        "Association, replication, and adjustment":"Statistical analysis",
        "Locked adjacent-section replication":"Validation",
        "Prior-task feasibility and participant-safe severity benchmark":"Secondary analysis",
      },
      "Results": {
        "Cohort and processing yield":"Cohort and image processing",
        "Medial discovery and lateral replication":"Discovery and replication",
        "Robustness and sensitivity":"Robustness and serial-section validation",
        "Biological component discrimination":"Biological interpretation",
        "Participant-safe severity benchmark":"Secondary severity analysis",
      },
    }
    for line in lines:
        if line.startswith("## Tables and figure legends"): skip_legends=True; continue
        if skip_legends and line.startswith("## Ethics and competing interests"): skip_legends=False
        if skip_legends or not line.strip(): continue
        if line=="## Status": skip_status=True; continue
        if skip_status and line=="## Abstract": skip_status=False
        if skip_status: continue
        if line.startswith("# ") and first:
            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); font(p.add_run(line[2:]),18,bold=True); first=False
            s=doc.add_paragraph("Research Article | Public human knee cartilage cohort"); s.paragraph_format.space_after=Pt(12); font(s.runs[0],9.5,italic=True,color=MUTED); continue
        if line.startswith("### "):
            label=line[4:]
            if section=="Abstract": continue
            replacement=retained_subheads.get(section,{}).get(label)
            if replacement: doc.add_heading(replacement,level=2)
            continue
        if line.startswith("## "):
            section=line[3:]; references=line=="## References"; doc.add_heading(section,level=1); continue
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; inline(p,line,9 if references else 10.5)
        if references: p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.0
        if "Whole-section processing yield is summarized in Table 1" in line: add_table(doc,1)
        if "reported in Table 2 and Figure 1" in line: add_table(doc,2); add_figure(doc,1)
        if "Figures 2–4" in line: add_figure(doc,2); add_figure(doc,3); add_figure(doc,4)
        if "Figures 5–8" in line: add_figure(doc,5); add_figure(doc,6); add_figure(doc,7); add_figure(doc,8)
    core=doc.core_properties; core.title="NOSTOS structural cartilage phenotype"; core.subject="Spatial-frequency histomorphology in human knee osteoarthritis"; core.author="NOSTOS study team"; core.comments="Tables and figures are embedded at first Results callout."
    doc.save(OUTPUT); print(OUTPUT)

if __name__=="__main__": build()
