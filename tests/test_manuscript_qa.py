from pathlib import Path

from docx import Document
from docx.shared import Pt
from PIL import Image

from nostos.validation.manuscript_qa import DEFAULT_REQUIRED_TEXT, build_manuscript_qa


def test_manuscript_qa_passes_a_complete_minimal_fixture(tmp_path: Path):
    source = tmp_path / "article.md"
    source.write_text("# Test article\n", encoding="utf-8")

    document = Document()
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(10)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(DEFAULT_REQUIRED_TEXT))
    run.font.name = "Times New Roman"
    docx_path = tmp_path / "article.docx"
    document.save(docx_path)

    render_dir = tmp_path / "render"
    render_dir.mkdir()
    Image.new("RGB", (100, 120), "white").save(render_dir / "page-1.png")
    pdf_path = render_dir / "article.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n" + b"0" * 100 + b"\n%%EOF\n")

    output = tmp_path / "qa.json"
    payload = build_manuscript_qa(
        project_root=tmp_path,
        manuscript_source=source,
        docx_path=docx_path,
        render_dir=render_dir,
        pdf_path=pdf_path,
        output_path=output,
        expected_pages=1,
        expected_media=0,
        visual_review_passed=True,
        visual_review_date="2026-08-28",
    )

    assert payload["status"] == "pass"
    assert payload["machine_status"] == "pass"
    assert payload["render"]["page_count"] == 1
    assert payload["nature_readiness"] == "not_ready"
    assert output.is_file()

